import os
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Brand Colors from brand.md (adapted from jdieselny genetics logo theme)
COLOR_MIDNIGHT = RGBColor(0x0C, 0x11, 0x17)   # #0C1117 (Primary dark accent)
COLOR_ICE_BLUE = RGBColor(0x8E, 0xCA, 0xE6)   # #8ECAE6 (Secondary light accent / borders)
COLOR_SLATE = RGBColor(0x47, 0x55, 0x69)      # #475569 (Muted slate)
COLOR_CHARCOAL = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A (Body Text)

COLOR_ICE_BLUE_HEX = "8ECAE6"
COLOR_LIGHT_BG_HEX = "F8FAFC"

def add_p_border_bottom(p, color_hex=COLOR_ICE_BLUE_HEX, size="12"):
    """Adds a structural horizontal border under the paragraph."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)  # size in 1/8 pt (12 = 1.5 pt)
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_p_border_left(p, color_hex=COLOR_ICE_BLUE_HEX, size="24"):
    """Adds a thick accent vertical border to the left of the paragraph."""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)  # 24 = 3 pt
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

def add_p_shading(p, color_hex=COLOR_LIGHT_BG_HEX):
    """Sets background shading color for the paragraph."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def add_page_number(run):
    """Inserts a dynamic PAGE field into a Word run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_paragraph_runs(p, text, is_bold=False, is_italic=False, font_name="Segoe UI", font_size=11, color_rgb=COLOR_CHARCOAL):
    """Helper to parse bold markdown tags and append runs to a paragraph."""
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = color_rgb
        run.italic = is_italic
        
        # Odd indices represent text inside ** bold tags
        if idx % 2 == 1:
            run.bold = True
        else:
            run.bold = is_bold

def render_markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return False
        
    doc = docx.Document()
    
    # Configure 1-inch margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Setup Header and Footer (tabs used for alignment)
    section.different_first_page_header_footer = False
    
    # Header Setup
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "The Cognitive Hypervisor\t\tWORKING DRAFT"
    for run in header_p.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_SLATE
        
    # Footer Setup
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "ECR-WG Specification | github.com/jdieselny/ecr-wg\t\tPage "
    for run in footer_p.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_SLATE
    
    # Append the PAGE field to footer
    page_run = footer_p.add_run()
    page_run.font.name = "Segoe UI"
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = COLOR_SLATE
    add_page_number(page_run)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.splitlines()
    in_code_block = False
    code_text = []
    
    for line in lines:
        stripped = line.strip()
        
        # Handle YAML Frontmatter
        if stripped == "---":
            # Just skip frontmatter markers and blocks
            continue
        if any(stripped.startswith(k) for k in ["aft:", "registrant:", "generated_at:", "file_role:", "agent:", "issued:", "issuer:", "source spec:", "status:", "version:", "target:", "date:"]):
            continue
            
        # Toggle Code Block
        if stripped.startswith("```"):
            if in_code_block:
                # Add accumulated code text
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                add_p_shading(p, COLOR_LIGHT_BG_HEX)  # Ice/Light Gray background shading
                code_str = "\n".join(code_text)
                run = p.add_run(code_str)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = COLOR_SLATE
                in_code_block = False
                code_text = []
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_text.append(line)
            continue
            
        # Parse Headings
        if stripped.startswith("# "):
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            add_paragraph_runs(p, title_text, is_bold=True, font_size=18, color_rgb=COLOR_MIDNIGHT)
            add_p_border_bottom(p, COLOR_ICE_BLUE_HEX, size="12")  # Ice Blue bottom divider
            continue
            
        if stripped.startswith("## "):
            h2_text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            add_paragraph_runs(p, h2_text, is_bold=True, font_size=14, color_rgb=COLOR_MIDNIGHT)
            continue
            
        if stripped.startswith("### "):
            h3_text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            add_paragraph_runs(p, h3_text, is_bold=True, font_size=12, color_rgb=COLOR_SLATE)
            continue
            
        # Parse Blockquotes
        if stripped.startswith("> "):
            bq_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_paragraph_runs(p, bq_text, is_italic=True, font_size=10.5, color_rgb=COLOR_SLATE)
            add_p_border_left(p, COLOR_ICE_BLUE_HEX, size="24")  # Ice Blue left border bar
            add_p_shading(p, COLOR_LIGHT_BG_HEX)  # Ice/Light Gray background shading
            continue
            
        # Parse Bullet Lists
        if stripped.startswith("* ") or stripped.startswith("- "):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_paragraph_runs(p, bullet_text, font_size=11, color_rgb=COLOR_CHARCOAL)
            continue
            
        # Parse Numbered Lists
        if re.match(r'^\d+\.\s', stripped):
            num_text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_paragraph_runs(p, num_text, font_size=11, color_rgb=COLOR_CHARCOAL)
            continue
            
        # Empty Line
        if not stripped:
            continue
            
        # Regular Body Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_paragraph_runs(p, line, font_size=11, color_rgb=COLOR_CHARCOAL)
        
    doc.save(docx_path)
    print(f"Successfully rendered DOCX to: {docx_path}")
    return True

if __name__ == "__main__":
    import sys
    md_file = r"C:\Users\jkintzele\Documents\ecr-wg\thesis\cognitive_hypervisor_thesis.md"
    docx_file = r"C:\Users\jkintzele\Documents\ecr-wg\thesis\cognitive_hypervisor_thesis.docx"
    
    if len(sys.argv) > 2:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]
        
    render_markdown_to_docx(md_file, docx_file)
